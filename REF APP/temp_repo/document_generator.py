from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
from typing import Dict, Any, List
import io
import logging
from date_utils import DateUtils

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentGenerator:
    """Generates Word documents for tender processing system."""
    
    def __init__(self):
        self.date_utils = DateUtils()
    
    def set_table_borders(self, table):
        """Add borders to table."""
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # Border width
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        
        tbl.tblPr.append(tblBorders)
    
    def generate_comparative_statement_doc(self, work: Dict[str, Any], bidders: List[Dict[str, Any]]) -> bytes:
        """Generate comparative statement in Word format matching PWD layout."""
        
        doc = Document()
        
        # Sort bidders by bid amount
        sorted_bidders = sorted(bidders, key=lambda x: x['bid_amount'])
        
        # Get work details
        work_name = work['work_name']
        nit_number = work['nit_number']
        work_info = work['work_info']
        estimated_cost = float(work_info['estimated_cost'])
        earnest_money = work_info['earnest_money']
        time_completion = work_info['time_of_completion']
        
        # Parse date
        original_date = work_info['date']
        parsed_date = self.date_utils.parse_date(original_date)
        formatted_date = self.date_utils.format_display_date(parsed_date) if parsed_date else original_date
        
        # Office header
        header = doc.add_paragraph('OFFICE OF THE EXECUTIVE ENGINEER PWD ELECTRIC DIVISION, UDAIPUR')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].bold = True
        header.runs[0].font.size = Pt(12)
        
        # Title
        title = doc.add_paragraph('COMPARATIVE STATEMENT OF TENDERS')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(12)
        title.runs[0].underline = True
        
        doc.add_paragraph()  # Space
        
        # Work details table
        details_table = doc.add_table(rows=6, cols=4)
        details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(details_table)
        
        # Fill work details
        details_table.cell(0, 0).text = "Name of Work:"
        details_table.cell(0, 1).text = work_name
        details_table.cell(0, 1).paragraphs[0].runs[0].bold = True
        
        details_table.cell(1, 0).text = "NIT No.:"
        details_table.cell(1, 1).text = nit_number
        details_table.cell(1, 2).text = "Date"
        details_table.cell(1, 3).text = formatted_date
        
        details_table.cell(2, 0).text = "1. Estimated amount for item in NIT Rs.:"
        details_table.cell(2, 1).text = f"{estimated_cost:,.0f}"
        details_table.cell(2, 2).text = "Earnest Money @2% Rs."
        details_table.cell(2, 3).text = str(earnest_money)
        
        details_table.cell(3, 0).text = "2. Amount of tender recommended for Rs:"
        details_table.cell(3, 1).text = f"{sorted_bidders[0]['bid_amount']:,.0f}"
        details_table.cell(3, 2).text = "Time for Completion Months"
        details_table.cell(3, 3).text = str(time_completion)
        
        details_table.cell(4, 0).text = "3 Estimated amount of item not included in the tender Rs.:"
        details_table.cell(4, 1).text = "Nil."
        details_table.cell(4, 2).text = "Date of calling NIT"
        details_table.cell(4, 3).text = formatted_date
        
        details_table.cell(5, 0).text = "4. Contingencies and other provision included in the estimate Rs:"
        details_table.cell(5, 1).text = "As per rules"
        details_table.cell(5, 2).text = "Date of Receipt of Tender"
        details_table.cell(5, 3).text = formatted_date
        
        doc.add_paragraph()  # Space
        
        # Main comparison table
        main_table = doc.add_table(rows=len(sorted_bidders) + 3, cols=4)
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(main_table)
        
        # Headers
        main_table.cell(0, 0).text = "S.No"
        main_table.cell(0, 1).text = "Bidder Name"
        main_table.cell(0, 2).text = "Estimated Cost Rs."
        main_table.cell(0, 3).text = "Quoted Amount Rs."
        
        # Subheaders
        main_table.cell(1, 2).text = "Quoted Percentage"
        
        # Make headers bold
        for i in range(4):
            main_table.cell(0, i).paragraphs[0].runs[0].bold = True
            main_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add bidder data
        for i, bidder in enumerate(sorted_bidders):
            row = i + 2
            main_table.cell(row, 0).text = str(i + 1)
            main_table.cell(row, 1).text = bidder['name']
            main_table.cell(row, 2).text = f"{bidder['percentage']:+.2f} BELOW" if bidder['percentage'] < 0 else f"{bidder['percentage']:+.2f} ABOVE"
            main_table.cell(row, 3).text = f"{bidder['bid_amount']:,.0f}"
            
            # Center align
            for j in range(4):
                main_table.cell(row, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add lowest bidder info
        lowest_row = len(sorted_bidders) + 2
        main_table.cell(lowest_row, 0).text = "Lowest"
        main_table.cell(lowest_row, 1).text = sorted_bidders[0]['name']
        main_table.cell(lowest_row, 2).text = f"{sorted_bidders[0]['percentage']:+.2f} BELOW"
        main_table.cell(lowest_row, 3).text = f"{sorted_bidders[0]['bid_amount']:,.0f}"
        
        # Make lowest row bold
        for j in range(4):
            main_table.cell(lowest_row, j).paragraphs[0].runs[0].bold = True
            main_table.cell(lowest_row, j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Space
        
        # Summary text
        summary = doc.add_paragraph(f"The tender of the lowest bidder {sorted_bidders[0]['name']}, "
                                   f"Udaipur @ {sorted_bidders[0]['percentage']:+.2f}% BELOW amounting to Rs. {sorted_bidders[0]['bid_amount']:,.0f}/-- "
                                   f"In words Rupees: Six Lakh Twenty Eight Thousand Eight Hundred Sixty One Only.")
        
        # Signature table
        sig_table = doc.add_table(rows=1, cols=4)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(sig_table)
        
        sig_table.cell(0, 0).text = "AR"
        sig_table.cell(0, 1).text = "DA"
        sig_table.cell(0, 2).text = "TA"
        sig_table.cell(0, 3).text = "EE"
        
        for i in range(4):
            sig_table.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig_table.cell(0, i).paragraphs[0].runs[0].bold = True
        
        # Footer signature section
        doc.add_paragraph()
        footer_sig = doc.add_table(rows=1, cols=4)
        footer_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(footer_sig)
        
        footer_sig.cell(0, 0).text = "Auditor"
        footer_sig.cell(0, 1).text = "Divisional Accountant"
        footer_sig.cell(0, 2).text = "TA"
        footer_sig.cell(0, 3).text = "Executive Engineer"
        
        for i in range(4):
            footer_sig.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_sig.cell(0, i).paragraphs[0].runs[0].bold = True
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
    
    def generate_scrutiny_sheet_doc(self, work: Dict[str, Any], bidders: List[Dict[str, Any]]) -> bytes:
        """Generate scrutiny sheet in Word format matching PWD layout."""
        
        doc = Document()
        
        # Sort bidders by bid amount
        sorted_bidders = sorted(bidders, key=lambda x: x['bid_amount'])
        lowest_bidder = sorted_bidders[0]
        
        # Get work details
        work_name = work['work_name']
        nit_number = work['nit_number']
        work_info = work['work_info']
        estimated_cost = float(work_info['estimated_cost'])
        
        # Parse date
        original_date = work_info['date']
        parsed_date = self.date_utils.parse_date(original_date)
        formatted_date = self.date_utils.format_display_date(parsed_date) if parsed_date else original_date
        
        # Title
        title = doc.add_paragraph('Scrutiny Sheet of Tender')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(14)
        title.runs[0].underline = True
        
        doc.add_paragraph()  # Space
        
        # Main scrutiny table
        table = doc.add_table(rows=16, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_table_borders(table)
        
        # Fill table data based on your reference image
        scrutiny_data = [
            ["1", "Head of Account", "PWD Electric Works"],
            ["2", "Name of work\nJob No.", f"{work_name}\n{nit_number}"],
            ["3", "Reference of ADM. Sanction\nAmount in Rs.", f"As per administrative approval\nRs. {estimated_cost:.0f}/-"],
            ["4", "Reference of technical\nsanction with amount", f"As per technical sanction for Rs. {estimated_cost:.0f}/-"],
            ["5", "Date of calling NIT", formatted_date],
            ["6", "Date of receipt of tender", formatted_date],
            ["7", "No. of tender sold", str(len(bidders))],
            ["8", "No. of tender received", str(len(bidders))],
            ["9", "Allotment of fund during the\ncurrent financial year", "Adequate."],
            ["10", "Expenditure up to last bill", "Nil."],
            ["11", "Lowest rate quoted and\ncondition if any", f"{lowest_bidder['percentage']:+.1f}% BELOW. No Condition."],
            ["12", "Financial implication of\ncondition if any in tender", "Not Applicable."],
            ["13", "Name of lowest contractor", lowest_bidder['name']],
            ["14", "Authority competent to\nsanction the tender", "The Executive Engineer"],
            ["15", "Validity of Tender\nValid Upto Dated", "20 Days\n13-04-25"],
            ["16", "Remarks if any", "None."]
        ]
        
        for i, (serial, description, value) in enumerate(scrutiny_data):
            table.cell(i, 0).text = serial
            table.cell(i, 1).text = description
            table.cell(i, 2).text = value
            
            # Format cells
            table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.cell(i, 0).paragraphs[0].runs[0].bold = True
            table.cell(i, 1).paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Space
        
        # Auditor signature
        auditor = doc.add_paragraph("AUDITOR")
        auditor.alignment = WD_ALIGN_PARAGRAPH.CENTER
        auditor.runs[0].bold = True
        auditor.runs[0].font.size = Pt(12)
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
    
    def generate_letter_of_acceptance_doc(self, work: Dict[str, Any], bidders: List[Dict[str, Any]]) -> bytes:
        """Generate Letter of Acceptance in Word format."""
        
        doc = Document()
        
        # Sort bidders and get L1
        sorted_bidders = sorted(bidders, key=lambda x: x['bid_amount'])
        l1_bidder = sorted_bidders[0]
        
        # Get work details
        work_name = work['work_name']
        nit_number = work['nit_number']
        work_info = work['work_info']
        estimated_cost = float(work_info['estimated_cost'])
        
        # Parse date
        original_date = work_info['date']
        parsed_date = self.date_utils.parse_date(original_date)
        formatted_date = self.date_utils.format_display_date(parsed_date) if parsed_date else original_date
        
        # Calculate stipulated start date (current date + 1 day)
        start_date = self.date_utils.add_days(datetime.now(), 1)
        start_date_str = self.date_utils.format_display_date(start_date)
        
        # Office header
        header = doc.add_paragraph('OFFICE OF THE EXECUTIVE ENGINEER PWD ELECTRIC DIVISION UDAIPUR')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].bold = True
        header.runs[0].font.size = Pt(12)
        
        # Title
        title = doc.add_paragraph('LETTER OF ACCEPTANCE')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(12)
        title.runs[0].underline = True
        
        doc.add_paragraph()  # Space
        
        # Letter content
        doc.add_paragraph(f"To,")
        doc.add_paragraph(f"{l1_bidder['name']}")
        doc.add_paragraph(f"{l1_bidder.get('address', 'Address on file')}")
        doc.add_paragraph()
        
        doc.add_paragraph(f"Subject: Letter of Acceptance for {work_name}")
        doc.add_paragraph(f"Reference: NIT No. {nit_number} dated {formatted_date}")
        doc.add_paragraph()
        
        doc.add_paragraph("Sir,")
        doc.add_paragraph()
        
        content = doc.add_paragraph(f"With reference to your tender dated {formatted_date} for the above mentioned work, "
                                   f"I am pleased to inform you that your tender has been accepted for Rs. {l1_bidder['bid_amount']:,.0f}/- "
                                   f"({l1_bidder['percentage']:+.2f}% of estimated cost).")
        
        doc.add_paragraph()
        
        doc.add_paragraph("You are hereby directed to commence the work immediately and complete the same within "
                         "the stipulated period as per the terms and conditions of the contract.")
        
        doc.add_paragraph()
        doc.add_paragraph(f"The work should be commenced from {start_date_str}.")
        doc.add_paragraph()
        
        doc.add_paragraph("Please acknowledge receipt of this letter and submit the required security deposit "
                         "and other documents as per the contract agreement.")
        
        doc.add_paragraph()
        doc.add_paragraph("Yours faithfully,")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        signature = doc.add_paragraph("Executive Engineer")
        signature.runs[0].bold = True
        doc.add_paragraph("PWD Electric Division")
        doc.add_paragraph("Udaipur")
        doc.add_paragraph()
        doc.add_paragraph(f"Date: {self.date_utils.get_current_date()}")
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
    
    def generate_work_order_doc(self, work: Dict[str, Any], bidders: List[Dict[str, Any]]) -> bytes:
        """Generate Work Order in Word format."""
        
        doc = Document()
        
        # Sort bidders and get L1
        sorted_bidders = sorted(bidders, key=lambda x: x['bid_amount'])
        l1_bidder = sorted_bidders[0]
        
        # Get work details
        work_name = work['work_name']
        nit_number = work['nit_number']
        work_info = work['work_info']
        estimated_cost = float(work_info['estimated_cost'])
        time_completion = work_info['time_of_completion']
        
        # Parse date
        original_date = work_info['date']
        parsed_date = self.date_utils.parse_date(original_date)
        formatted_date = self.date_utils.format_display_date(parsed_date) if parsed_date else original_date
        
        # Calculate stipulated start date (current date + 1 day)
        start_date = self.date_utils.add_days(datetime.now(), 1)
        start_date_str = self.date_utils.format_display_date(start_date)
        
        # Office header
        header = doc.add_paragraph('OFFICE OF THE EXECUTIVE ENGINEER PWD ELECTRIC DIVISION UDAIPUR')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].bold = True
        header.runs[0].font.size = Pt(12)
        
        # Title
        title = doc.add_paragraph('WORK ORDER')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(12)
        title.runs[0].underline = True
        
        doc.add_paragraph()  # Space
        
        # Work order details
        doc.add_paragraph(f"Work Order No.: WO/{nit_number}/{datetime.now().year}")
        doc.add_paragraph(f"Date: {self.date_utils.get_current_date()}")
        doc.add_paragraph()
        
        doc.add_paragraph(f"To,")
        doc.add_paragraph(f"{l1_bidder['name']}")
        doc.add_paragraph(f"{l1_bidder.get('address', 'Address on file')}")
        doc.add_paragraph()
        
        doc.add_paragraph(f"Subject: Work Order for {work_name}")
        doc.add_paragraph(f"Reference: NIT No. {nit_number} dated {formatted_date}")
        doc.add_paragraph()
        
        doc.add_paragraph("Sir,")
        doc.add_paragraph()
        
        doc.add_paragraph("You are hereby directed to execute the following work as per the terms and conditions "
                         "of the contract:")
        doc.add_paragraph()
        
        work_details = doc.add_paragraph("Work Details:")
        work_details.runs[0].bold = True
        doc.add_paragraph(f"Name of Work: {work_name}")
        doc.add_paragraph(f"Contract Amount: Rs. {l1_bidder['bid_amount']:,.0f}/-")
        doc.add_paragraph(f"Time of Completion: {time_completion}")
        doc.add_paragraph(f"Stipulated Date of Start: {start_date_str}")
        doc.add_paragraph(f"Earnest Money: Rs. {l1_bidder['earnest_money']}")
        doc.add_paragraph()
        
        doc.add_paragraph("You are directed to commence the work immediately and complete the same within "
                         "the stipulated time as per the agreement.")
        doc.add_paragraph()
        
        doc.add_paragraph("All terms and conditions as per the tender document and contract agreement "
                         "shall be applicable.")
        doc.add_paragraph()
        
        doc.add_paragraph("This work order is issued subject to the fulfillment of all contractual "
                         "obligations including submission of required security deposit.")
        doc.add_paragraph()
        
        doc.add_paragraph("Yours faithfully,")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        signature = doc.add_paragraph("Executive Engineer")
        signature.runs[0].bold = True
        doc.add_paragraph("PWD Electric Division")
        doc.add_paragraph("Udaipur")
        
        # Save to bytes
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer.getvalue()
